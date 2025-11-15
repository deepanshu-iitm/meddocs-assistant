"""
Document processing service for extracting content from various file types
"""
import os
import io
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import pytesseract
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
import json
import base64

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing different types of medical documents"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process a document and extract content, tables, and metadata
        
        Args:
            file_path: Path to the document file
            file_type: Type of document (pdf, docx, xlsx, image)
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            if file_type.lower() == 'pdf':
                return self._process_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._process_docx(file_path)
            elif file_type.lower() in ['xlsx', 'xls']:
                return self._process_excel(file_path)
            elif file_type.lower() in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                return self._process_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files and extract text, tables, and images"""
        content = {
            'text': '',
            'pages': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            # Use pdfplumber for better table extraction
            with pdfplumber.open(file_path) as pdf:
                content['metadata'] = {
                    'num_pages': len(pdf.pages),
                    'title': pdf.metadata.get('Title', ''),
                    'author': pdf.metadata.get('Author', ''),
                    'subject': pdf.metadata.get('Subject', ''),
                    'creator': pdf.metadata.get('Creator', '')
                }
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    content['text'] += page_text + "\n\n"
                    
                    # Store page-specific content
                    page_content = {
                        'page_number': page_num,
                        'text': page_text,
                        'tables': [],
                        'images': []
                    }
                    
                    # Extract tables from page
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_data = {
                                'page_number': page_num,
                                'table_index': table_idx,
                                'data': table,
                                'text_representation': self._table_to_text(table)
                            }
                            content['tables'].append(table_data)
                            page_content['tables'].append(table_data)
                    
                    content['pages'].append(page_content)
            
            # Also try PyPDF2 for additional metadata
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    if pdf_reader.metadata:
                        content['metadata'].update({
                            'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                            'modification_date': str(pdf_reader.metadata.get('/ModDate', '')),
                            'producer': str(pdf_reader.metadata.get('/Producer', ''))
                        })
            except Exception as e:
                logger.warning(f"Could not extract additional PDF metadata: {e}")
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
        
        return content
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents and extract text and tables"""
        content = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            content['metadata'] = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    content['text'] += para.text + "\n"
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                table_info = {
                    'table_index': table_idx,
                    'data': table_data,
                    'text_representation': self._table_to_text(table_data)
                }
                content['tables'].append(table_info)
                content['text'] += f"\n\nTable {table_idx + 1}:\n{table_info['text_representation']}\n"
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
        
        return content
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel files and extract data from sheets"""
        content = {
            'text': '',
            'sheets': [],
            'tables': [],
            'metadata': {}
        }
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Extract metadata
            content['metadata'] = {
                'sheet_names': workbook.sheetnames,
                'num_sheets': len(workbook.sheetnames),
                'title': workbook.properties.title or '',
                'creator': workbook.properties.creator or '',
                'created': str(workbook.properties.created) if workbook.properties.created else ''
            }
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # Get all data from sheet
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):  # Skip empty rows
                        sheet_data.append([str(cell) if cell is not None else '' for cell in row])
                
                if sheet_data:
                    sheet_info = {
                        'sheet_name': sheet_name,
                        'data': sheet_data,
                        'text_representation': self._table_to_text(sheet_data)
                    }
                    content['sheets'].append(sheet_info)
                    content['tables'].append({
                        'sheet_name': sheet_name,
                        'data': sheet_data,
                        'text_representation': sheet_info['text_representation']
                    })
                    content['text'] += f"\n\nSheet: {sheet_name}\n{sheet_info['text_representation']}\n"
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            raise
        
        return content
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process images using OCR to extract text"""
        content = {
            'text': '',
            'ocr_text': '',
            'metadata': {}
        }
        
        try:
            # Open image
            image = Image.open(file_path)
            
            # Extract metadata
            content['metadata'] = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            }
            
            # Perform OCR
            try:
                ocr_text = pytesseract.image_to_string(image, config='--psm 6')
                content['ocr_text'] = ocr_text
                content['text'] = ocr_text
            except Exception as e:
                logger.warning(f"OCR failed for image {file_path}: {e}")
                content['ocr_text'] = ""
                content['text'] = ""
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise
        
        return content
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable text format"""
        if not table_data:
            return ""
        
        # Create text representation of table
        text_lines = []
        for row in table_data:
            # Clean and join cells
            clean_row = [str(cell).strip() for cell in row if cell is not None]
            if clean_row:  # Only add non-empty rows
                text_lines.append(" | ".join(clean_row))
        
        return "\n".join(text_lines)
    
    def chunk_content(self, content: str, metadata: Optional[Dict] = None) -> List[Dict[str, Any]]:
        """
        Split content into chunks for vector storage
        
        Args:
            content: Text content to chunk
            metadata: Additional metadata to include with chunks
            
        Returns:
            List of chunk dictionaries
        """
        if not content.strip():
            return []
        
        try:
            # Split text into chunks
            chunks = self.text_splitter.split_text(content)
            
            chunk_list = []
            for i, chunk in enumerate(chunks):
                chunk_data = {
                    'content': chunk,
                    'chunk_index': i,
                    'chunk_type': 'text',
                    'metadata': metadata or {}
                }
                chunk_list.append(chunk_data)
            
            return chunk_list
            
        except Exception as e:
            logger.error(f"Error chunking content: {e}")
            raise
    
    def extract_medical_sections(self, content: str) -> Dict[str, str]:
        """
        Extract common medical document sections
        
        Args:
            content: Full document text
            
        Returns:
            Dictionary mapping section names to content
        """
        sections = {}
        
        # Common medical document section patterns
        section_patterns = {
            'introduction': [r'introduction', r'background', r'overview'],
            'clinical_findings': [r'clinical findings', r'findings', r'results', r'observations'],
            'patient_data': [r'patient data', r'patient information', r'demographics'],
            'diagnosis': [r'diagnosis', r'diagnostic', r'assessment'],
            'treatment': [r'treatment', r'therapy', r'intervention', r'medication'],
            'summary': [r'summary', r'conclusion', r'conclusions'],
            'recommendations': [r'recommendations', r'recommendations', r'next steps']
        }
        
        # Simple section extraction based on headers
        lines = content.split('\n')
        current_section = 'general'
        sections[current_section] = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            section_found = False
            for section_name, patterns in section_patterns.items():
                for pattern in patterns:
                    if pattern in line_lower and len(line.strip()) < 100:  # Likely a header
                        current_section = section_name
                        if current_section not in sections:
                            sections[current_section] = []
                        section_found = True
                        break
                if section_found:
                    break
            
            if not section_found and line.strip():
                sections[current_section].append(line)
        
        # Convert lists to strings
        for section_name in sections:
            sections[section_name] = '\n'.join(sections[section_name]).strip()
        
        # Remove empty sections
        sections = {k: v for k, v in sections.items() if v.strip()}
        
        return sections
